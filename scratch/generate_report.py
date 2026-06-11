import json
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    """Sets background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def create_report():
    # Paths
    json_path = r"d:/Xu Ly Anh So/ProjectNhom/vehicle-type-recognition/outputs/evaluation_yolo_cls_best.json"
    docx_path = r"d:/Xu Ly Anh So/ProjectNhom/vehicle-type-recognition/docs/YOLO_TRAINING_REPORT.docx"
    
    # Load JSON
    if os.path.exists(json_path):
        with open(json_path, 'r', encoding='utf-8') as f:
            eval_data = json.load(f)
    else:
        eval_data = {}

    doc = Document()
    
    # Page Setup
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Styles Setup
    style_normal = doc.styles['Normal']
    font_normal = style_normal.font
    font_normal.name = 'Arial'
    font_normal.size = Pt(11)
    font_normal.color.rgb = RGBColor(0x2C, 0x3E, 0x50) # Dark gray/blue text

    # Helper function to add headings with specific color and size
    def add_heading_1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x1B, 0x4F, 0x30) # Dark green
        return p

    def add_heading_2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32) # Medium green
        return p

    def add_heading_3(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x42, 0x42, 0x42) # Gray
        return p

    # --- TITLE ---
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("BÁO CÁO TOÀN DIỆN HUẤN LUYỆN, ĐÁNH GIÁ VÀ TÍCH HỢP\nMÔ HÌNH PHÂN LOẠI YOLOv8-CLS")
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = RGBColor(0x1B, 0x4F, 0x30)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("Hệ Thống Nhận Diện Loại Phương Tiện Giao Thông (Vehicle Type Recognition)")
    sub_run.italic = True
    sub_run.font.size = Pt(12)
    sub_run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
    
    doc.add_paragraph() # Spacer

    # --- CHƯƠNG 1 ---
    add_heading_1("1. KIẾN TRÚC HỆ THỐNG VÀ CHIẾN LƯỢC DỮ LIỆU (DATA FLOWS)")
    
    p_intro = doc.add_paragraph(
        "Báo cáo này cung cấp cái nhìn toàn diện về quy trình xử lý dữ liệu (data flow), huấn luyện và đánh giá mô hình "
        "YOLOv8-cls trong dự án Nhận Diện Loại Phương Tiện Giao Thông (Vehicle Type Recognition). Hệ thống tích hợp song song "
        "ba mô hình: ResNet-50 (CNN baseline), Vision Transformer (ViT - global context), và YOLO-cls (fast inference)."
    )

    add_heading_2("1.1. Thống Kê Phân Bổ Tập Dữ Liệu Gốc (Vehicle-10)")
    doc.add_paragraph(
        "Bộ dữ liệu Vehicle-10 chứa tổng cộng 36,006 ảnh RGB được phân chia thành 10 lớp phương tiện giao thông. "
        "Dữ liệu được phân chia thành các nhóm kích thước lớp khác nhau (Minority, Medium, Large) nhằm chuẩn bị cho chiến lược cân bằng:"
    )

    # Table 1: Dataset Distribution
    table_ds = doc.add_table(rows=11, cols=3)
    table_ds.style = 'Table Grid'
    headers_ds = ["Lớp (Class)", "Tên Tiếng Việt", "Phân Nhóm Kích Thước (Group)"]
    for i, h in enumerate(headers_ds):
        cell = table_ds.cell(0, i)
        cell.text = h
        set_cell_background(cell, "225F47")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].runs[0].font.bold = True

    ds_rows = [
        ("bicycle", "Xe đạp", "Minority (Thiểu số)"),
        ("boat", "Thuyền/Tàu thủy", "Large (Quy mô lớn)"),
        ("bus", "Xe buýt", "Medium (Trung bình)"),
        ("car", "Ô tô con", "Large (Quy mô lớn)"),
        ("helicopter", "Trực thăng", "Minority (Thiểu số)"),
        ("minibus", "Xe khách nhỏ/16 chỗ", "Minority (Thiểu số)"),
        ("motorcycle", "Xe máy/Mô tô", "Medium (Trung bình)"),
        ("taxi", "Xe Taxi", "Minority (Thiểu số)"),
        ("train", "Tàu hỏa", "Minority (Thiểu số)"),
        ("truck", "Xe tải", "Medium (Trung bình)")
    ]
    for r_idx, (cls, vn, grp) in enumerate(ds_rows, start=1):
        table_ds.cell(r_idx, 0).text = cls
        table_ds.cell(r_idx, 1).text = vn
        table_ds.cell(r_idx, 2).text = grp
        if "Minority" in grp:
            set_cell_background(table_ds.cell(r_idx, 2), "FEF9E7") # Light yellow alert for minority

    add_heading_2("1.2. Chiến Lược Phân Chia Tập Dữ Liệu (Split Strategy)")
    doc.add_paragraph(
        "Để đảm bảo đánh giá khách quan và kiểm tra khả năng tổng quát hóa (generalization), bộ dữ liệu được chia theo tỷ lệ nghiêm ngặt "
        "và độc lập như sau:"
    )

    # Table 2: Splits
    table_splits = doc.add_table(rows=5, cols=4)
    table_splits.style = 'Table Grid'
    headers_split = ["Tập Dữ Liệu (Split)", "Tỷ Lệ (%)", "Số Lượng Ảnh Xấp Xỉ", "Vai Trò Trong Hệ Thống"]
    for i, h in enumerate(headers_split):
        cell = table_splits.cell(0, i)
        cell.text = h
        set_cell_background(cell, "225F47")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].runs[0].font.bold = True

    split_rows = [
        ("train", "85%", "~30,605", "Dùng để huấn luyện; tập duy nhất được áp dụng cân bằng và tăng cường."),
        ("valid_unseen", "5%", "~1,800", "Tập validation độc lập dùng để đánh giá hiệu năng trực tiếp khi huấn luyện."),
        ("test", "10%", "~3,601", "Tập kiểm thử độc lập cuối cùng, chỉ dùng sau khi đã đóng băng checkpoint."),
        ("valid_traincopy", "N/A (Copy từ train)", "~1,800", "Tập tham chiếu phụ trợ theo yêu cầu giáo trình (không đại diện cho Generalization).")
    ]
    for r_idx, (name, pct, cnt, role) in enumerate(split_rows, start=1):
        table_splits.cell(r_idx, 0).text = name
        table_splits.cell(r_idx, 1).text = pct
        table_splits.cell(r_idx, 2).text = cnt
        table_splits.cell(r_idx, 3).text = role

    # --- CHƯƠNG 2 ---
    add_heading_1("2. QUY TRÌNH TIỀN XỬ LÝ VÀ TĂNG CƯỜNG DỮ LIỆU (PREPROCESSING & AUGMENTATION)")
    
    add_heading_2("2.1. Quy Trình Cân Bằng Lớp (Class Balancing Flow)")
    doc.add_paragraph(
        "Nhằm giải quyết hiện tượng mất cân bằng dữ liệu giữa các lớp lớn (Boat, Car) và các lớp thiểu số (Taxi, Train, Helicopter), "
        "quy trình cân bằng lớp chỉ áp dụng trên tập huấn luyện (train split). Quá trình chuẩn bị được thiết lập như sau:\n"
        "• Đối với lớp quy mô lớn (Boat, Car): Áp dụng chính sách cắt giảm (under-generation) và phân bổ cố định, giữ 70% ở điều kiện bình thường và 30% ở các môi trường thời tiết khác nhau.\n"
        "• Đối với lớp thiểu số (Helicopter, Taxi, Train, Bicycle, Minibus): Tạo thêm các biến thể hình học (quay ảnh, lật ngang, dịch chuyển, thay đổi phối cảnh) và biến thể độ tương phản cho đến khi đạt tối thiểu 70% hạn ngạch (quota) của lớp lớn nhất."
    )

    add_heading_2("2.2. Quy Trình Tăng Cường Ngoại Tuyến (Offline Augmentation Flow)")
    doc.add_paragraph(
        "Ảnh sau khi cân bằng được đưa vào quy trình mô phỏng môi trường thời tiết thực tế theo tỷ lệ cố định:"
    )

    # Table 3: Augmentations
    table_aug = doc.add_table(rows=5, cols=3)
    table_aug.style = 'Table Grid'
    headers_aug = ["Nhóm Môi Trường", "Tỷ Lệ Phân Bổ", "Mô Tả Biến Đổi Vật Lý"]
    for i, h in enumerate(headers_aug):
        cell = table_aug.cell(0, i)
        cell.text = h
        set_cell_background(cell, "225F47")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].runs[0].font.bold = True

    aug_rows = [
        ("normal", "70%", "Ảnh gốc đi qua bộ tiền xử lý chuẩn Base Pipeline."),
        ("rain", "10%", "Giả lập vệt mưa rơi, hiệu ứng hạt nước xước mờ dọc màn hình và giảm nhẹ độ bão hòa màu."),
        ("sun", "10%", "Giả lập hiệu ứng chói lóa ánh nắng mặt trời (lens flare) và tăng độ tương phản vùng sáng."),
        ("night", "10%", "Giả lập môi trường thiếu sáng, giảm độ sáng toàn cục và tăng nhiễu hạt (ISO noise).")
    ]
    for r_idx, (name, pct, desc) in enumerate(aug_rows, start=1):
        table_aug.cell(r_idx, 0).text = name
        table_aug.cell(r_idx, 1).text = pct
        table_aug.cell(r_idx, 2).text = desc

    add_heading_2("2.3. Bộ Tiền Xử Lý Chuẩn (Base Pipeline Preprocessing)")
    doc.add_paragraph(
        "Tất cả các tập dữ liệu đánh giá (valid_unseen, test) và ảnh đầu vào khi suy luận thực tế (deployment) "
        "đều đi qua một đường ống tiền xử lý duy nhất (Base Pipeline):\n"
        "1. Thay đổi kích thước ảnh nhưng vẫn giữ nguyên tỷ lệ khung hình gốc (preserve aspect ratio) để tránh biến dạng hình học của xe.\n"
        "2. Đệm viền màu đen (zero-padding) xung quanh phần thiếu để đưa ảnh về kích thước vuông chuẩn 224x224 pixel.\n"
        "Không có bất kỳ bộ lọc làm sạch hay kỹ thuật làm nhiễu vật lý nào khác được áp dụng trong Base Pipeline nhằm đảm bảo tính khách quan."
    )

    # --- CHƯƠNG 3 ---
    add_heading_1("3. ĐẶC THÙ TÍCH HỢP MÔ HÌNH PHÂN LOẠI YOLO-CLS")
    doc.add_paragraph(
        "Khác biệt lớn nhất giữa mô hình YOLO-cls (Ultralytics) và các kiến trúc PyTorch tiêu chuẩn (ResNet-50, ViT) nằm ở cách thức nạp trọng số, "
        "đầu vào và cấu trúc trả về kết quả:"
    )

    # Table 4: YOLO vs PyTorch
    table_diff = doc.add_table(rows=5, cols=3)
    table_diff.style = 'Table Grid'
    headers_diff = ["Đặc Tính Xử Lý", "Mô Hình PyTorch Tiêu Chuẩn (ResNet-50, ViT)", "Mô Hình YOLO-cls (Ultralytics)"]
    for i, h in enumerate(headers_diff):
        cell = table_diff.cell(0, i)
        cell.text = h
        set_cell_background(cell, "225F47")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].runs[0].font.bold = True

    diff_rows = [
        ("Định dạng file lưu trữ", ".pth (chỉ lưu các tham số trọng số state_dict)", ".pt (lưu toàn bộ đồ thị mạng, cấu trúc và siêu tham số huấn luyện)"),
        ("Phương thức nạp mô hình", "torch.load() -> khởi tạo Class PyTorch -> load_state_dict()", "import YOLO -> model = YOLO('best.pt')"),
        ("Đầu vào suy luận", "Bắt buộc Tensor PyTorch chuẩn hóa dạng (1, 3, 224, 224)", "Đa dạng: Ảnh PIL trực tiếp, NumPy array, đường dẫn ảnh, hoặc Tensor"),
        ("Đầu ra dự đoán", "Logits thô (chưa qua chuẩn hóa, cần Softmax thủ công)", "Danh sách đối tượng Results chứa tensor xác suất Probs sẵn sàng")
    ]
    for r_idx, (name, torch_desc, yolo_desc) in enumerate(diff_rows, start=1):
        table_diff.cell(r_idx, 0).text = name
        table_diff.cell(r_idx, 1).text = torch_desc
        table_diff.cell(r_idx, 2).text = yolo_desc

    # --- CHƯƠNG 4 ---
    add_heading_1("4. KẾT QUẢ HUÂN LUYỆN VÀ PHÂN TÍCH ĐƯỜNG CONG HỘI TỤ")
    
    add_heading_2("4.1. Phân Tích Đường Cong Giảm Loss và Accuracy")
    doc.add_paragraph(
        "Quá trình huấn luyện YOLOv8-cls được theo dõi qua 25 Epochs thực tế:\n"
        "• Giai đoạn đầu (Epoch 1 - 3): Tốc độ học cực kỳ nhanh. Train Loss giảm mạnh từ khoảng 1.20 xuống dưới 0.45. Điều này cho thấy mô hình thích ứng nhanh với các đặc trưng lớn của xe.\n"
        "• Giai đoạn bão hòa (Epoch 10 - 25): Train Loss tiếp tục giảm mịn dần và tiến về mức sát 0.10 ở Epoch 25. Tuy nhiên, Val Loss bắt đầu đi ngang và dao động quanh ngưỡng 0.25 - 0.28.\n"
        "• Hiện tượng quá khớp nhẹ (Early Overfitting): Do đường Train Loss tiếp tục đi xuống sâu nhưng đường Val Loss bị bão hòa phẳng, đây là biểu hiện của việc mô hình bắt đầu học thuộc lòng chi tiết nhỏ của tập Train. Tuy nhiên, hiện tượng này ở mức chấp nhận được vì Val Loss không bị uốn cong vọt lên lại."
    )

    if eval_data:
        headers = ["Tập dữ liệu (Split)", "Độ chính xác (Accuracy)", "Macro F1-Score", "Weighted F1-Score", "Số lượng mẫu (Samples)"]
        splits = ["valid_unseen", "test", "valid_traincopy"]
        split_labels = {
            "valid_unseen": "Valid Unseen (Độc lập)",
            "test": "Test (Đánh giá chính thức)",
            "valid_traincopy": "Valid Traincopy (Tham chiếu)"
        }
        classes = eval_data.get("class_names", [])
        headers_cls = ["Lớp (Class)", "Độ chính xác dự đoán đúng (Precision)", "Tỷ lệ không bỏ sót (Recall)", "F1-Score (Tổng hợp)", "Số lượng mẫu (Support)"]

        add_heading_2("4.2. Hiệu Năng So Sánh Giữa Các Tập Dữ Liệu")
        doc.add_paragraph(
            "Bảng dưới đây tổng hợp kết quả đánh giá thực tế của mô hình YOLOv8-cls tại checkpoint tốt nhất:"
        )

        table_perf = doc.add_table(rows=4, cols=5)
        table_perf.style = 'Table Grid'
        for i, h in enumerate(headers):
            cell = table_perf.cell(0, i)
            cell.text = h
            set_cell_background(cell, "225F47")
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            cell.paragraphs[0].runs[0].font.bold = True

        for row_idx, split_key in enumerate(splits, start=1):
            s_data = eval_data[split_key]
            s_rep = s_data["classification_report"]
            table_perf.cell(row_idx, 0).text = split_labels[split_key]
            table_perf.cell(row_idx, 1).text = f"{s_data['accuracy']*100:.2f}%"
            table_perf.cell(row_idx, 2).text = f"{s_rep['macro avg']['f1-score']*100:.2f}%"
            table_perf.cell(row_idx, 3).text = f"{s_rep['weighted avg']['f1-score']*100:.2f}%"
            table_perf.cell(row_idx, 4).text = f"{s_data['samples']}"
            set_cell_background(table_perf.cell(row_idx, 1), "F4F9F6")

        # --- CHƯƠNG 5 ---
        add_heading_1("5. BÁO CÁO PHÂN LỚP CHI TIẾT TỪNG LỚP VÀ PHÂN TÍCH SAI SỐ (TEST SPLIT)")
        
        add_heading_2("5.1. Bảng Chỉ Số Phân Lớp Lớp Chi Tiết (Classification Report)")
        doc.add_paragraph(
            "Đánh giá chi tiết hiệu năng phân lớp của YOLOv8-cls trên tập kiểm thử chính thức (Test Split - 3,601 ảnh) như sau:"
        )

        table_test_cls = doc.add_table(rows=len(classes) + 2, cols=5)
        table_test_cls.style = 'Table Grid'
        for i, h in enumerate(headers_cls):
            cell = table_test_cls.cell(0, i)
            cell.text = h
            set_cell_background(cell, "225F47")
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            cell.paragraphs[0].runs[0].font.bold = True

        test_report = eval_data["test"]["classification_report"]
        for row_idx, cls in enumerate(classes, start=1):
            cls_m = test_report[cls]
            table_test_cls.cell(row_idx, 0).text = cls
            table_test_cls.cell(row_idx, 1).text = f"{cls_m['precision']*100:.1f}%"
            table_test_cls.cell(row_idx, 2).text = f"{cls_m['recall']*100:.1f}%"
            table_test_cls.cell(row_idx, 3).text = f"{cls_m['f1-score']*100:.1f}%"
            table_test_cls.cell(row_idx, 4).text = f"{int(cls_m['support'])}"
            
            # Highlight poor classes
            if cls_m['f1-score'] < 0.82:
                set_cell_background(table_test_cls.cell(row_idx, 3), "FADBD8")
                set_cell_background(table_test_cls.cell(row_idx, 0), "FDEDEC")

        macro_m = test_report["macro avg"]
        last_idx = len(classes) + 1
        table_test_cls.cell(last_idx, 0).text = "Macro Average"
        table_test_cls.cell(last_idx, 1).text = f"{macro_m['precision']*100:.1f}%"
        table_test_cls.cell(last_idx, 2).text = f"{macro_m['recall']*100:.1f}%"
        table_test_cls.cell(last_idx, 3).text = f"{macro_m['f1-score']*100:.1f}%"
        table_test_cls.cell(last_idx, 4).text = f"{int(macro_m['support'])}"
        set_cell_background(table_test_cls.cell(last_idx, 0), "EAECEE")
        set_cell_background(table_test_cls.cell(last_idx, 3), "D5F5E3")
        for i in range(5):
            table_test_cls.cell(last_idx, i).paragraphs[0].runs[0].font.bold = True

        add_heading_2("5.2. Phân Tích Ma Trận Nhầm Lẫn (Confusion Matrix)")
        doc.add_paragraph(
            "Phân tích sâu ma trận nhầm lẫn cho thấy các lỗi phân loại hệ thống nghiêm trọng nhất:\n"
            "• Nhầm lẫn Taxi -> Car (Lỗi Recall lớn nhất): Lớp Taxi có recall chỉ đạt 63.7%. Có tới 33 trong tổng số 91 ảnh taxi bị dự đoán nhầm thành ô tô con (Car). Điều này chứng tỏ mô hình gặp khó khăn lớn trong việc phát hiện mào taxi nhỏ trên nóc xe.\n"
            "• Nhầm lẫn Truck -> Car & Bus: 37 xe tải bị nhầm thành ô tô con, 15 xe tải nhầm thành xe buýt.\n"
            "• Nhầm lẫn Minibus -> Car & Truck: Xe minibus 16 chỗ bị dự đoán sai thành 12 ảnh Car và 12 ảnh Truck do phom dáng thon dài trung gian của dòng xe này.\n"
            "• Nhầm lẫn Bus -> Boat (Ảnh phản chiếu ánh sáng): Dưới điều kiện ánh sáng xanh dương cabin phản chiếu mạnh từ kính xe buýt ban đêm, mô hình bị lừa nhận dạng thành thuyền (boat) do hình học bo tròn và màu xanh phản chiếu đặc trưng của khoang tàu thủy."
        )

    # --- CHƯƠNG 6 ---
    add_heading_1("6. ĐỀ XUẤT PHƯƠNG ÁN TĂNG ĐỘ CHÍNH XÁC (IMPROVEMENT STRATEGIES)")
    doc.add_paragraph(
        "Nhằm nâng cao độ chính xác của hệ thống, đặc biệt là khắc phục các lỗi phân loại ở Taxi, Minibus và Truck, "
        "bốn phương án kỹ thuật sau được đề xuất áp dụng cho chu kỳ huấn luyện tiếp theo:"
    )

    doc.add_paragraph(
        "1. Huấn luyện bằng Weighted Dataloader (Bộ tải dữ liệu có trọng số):\n"
        "Thay vì nạp ảnh ngẫu nhiên, sử dụng WeightedRandomSampler để tăng tần suất xuất hiện của các ảnh Taxi (91 ảnh) và Minibus (148 ảnh) "
        "trong mỗi batch huấn luyện. Kỹ thuật này giúp mô hình học đều đặn các lớp thiểu số mà không gây hiện tượng quá khớp tĩnh."
    )

    doc.add_paragraph(
        "2. Tăng kích thước ảnh huấn luyện đầu vào (imgsz = 320 hoặc 640):\n"
        "Việc tăng kích thước ảnh giúp bảo toàn chi tiết nhỏ của vật thể (như chữ TAXI dán trên cửa hoặc mào taxi trên nóc). "
        "Nhờ đó, mạng tích chập sâu của YOLO trích xuất được các đặc trưng tốt hơn."
    )

    doc.add_paragraph(
        "3. Tinh chỉnh Tăng cường dữ liệu động (Dynamic Augmentation):\n"
        "Áp dụng thêm các kỹ thuật crop ngẫu nhiên (Random Crop) tập trung nửa trên cabin xe để học đặc trưng mào xe taxi. "
        "Đồng thời, áp dụng Random Erasing để mô phỏng vật thể bị che khuất một phần."
    )

    doc.add_paragraph(
        "4. Nâng cấp phiên bản mô hình (Scale Up Model):\n"
        "Chuyển đổi từ YOLOv8n-cls (Nano, 1.45M tham số) lên YOLOv8s-cls (Small, ~5M tham số) hoặc YOLOv8m-cls (Medium, ~11M tham số). "
        "Mô hình lớn hơn sẽ có khả năng ghi nhớ đặc trưng tốt hơn và tăng độ chính xác phân biệt các loại xe tương đồng."
    )

    doc.save(docx_path)
    print(f"Word report successfully saved to: {docx_path}")

if __name__ == "__main__":
    create_report()
